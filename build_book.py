# -*- coding: utf-8 -*-
"""整合《百年回响，次元归途》全书，导出整合 Markdown / Word(.docx) / EPUB(.epub)。

使用说明：
1. 将底图放入 BASE/cover_art/，命名为 cover_bg.png（可用 image_gen 生成）
2. 运行：python build_book.py
3. 输出：百年回响，次元归途（全本）.{md,docx,epub} + cover_art/cover_final.png
"""
import os
import re

BASE = r"d:\Documents\AI_novel\ai_novel\novels\百年回响，次元归途"
CHAPTERS = os.path.join(BASE, "chapters")
COVER_DIR = os.path.join(BASE, "cover_art")
COVER_BG = os.path.join(COVER_DIR, "cover_bg.png")
COVER_FINAL = os.path.join(COVER_DIR, "cover_final.png")

BOOK_TITLE = "百年回响，次元归途"
SUBTITLE = "次元有壁，爱意无疆。百年回响，终有归途。"
AUTHOR = "花小香香"

# 章节文件顺序与卷归属
ORDER = [
    "第01章-一九二六，烟雨沪上.md",
    "第02章-乱世心事，尽付虚空.md",
    "第03章-单向时空，百年遥望.md",
    "第04章-乱世余生，遗憾落幕.md",
    "第05章-百年沉淀，医生沈聿.md",
    "第06章-盛世重生，少女林晚.md",
    "第07章-次元破壁，双向互通.md",
    "第08章-第一声回应，百年如愿.md",
    "第09章-独属于两人的秘密.md",
    "第10章-双向治愈，彼此救赎.md",
    "第11章-次元恋爱的极致拉扯.md",
    "第12章-七夕星河，时空共赴.md",
    "第13章-执念成研，破壁之路.md",
    "第14章-百年破壁，终得相逢.md",
    "第15章-次元归期，余生相守.md",
    "番外一-雨夜微光.md",
    "番外二-少年心事.md",
    "番外三-归期.md",
    "番外四-命中注定.md",
]

VOLUMES = [
    ("第一卷 梧桐旧梦・百年无声", 1, 4),
    ("第二卷 盛世重逢・次元共鸣", 5, 8),
    ("第三卷 次元相恋・无人知晓的深爱", 9, 11),
    ("第四卷 星河七夕・百年一遇", 12, 12),
    ("第五卷 破壁终章・岁岁归途", 13, 15),
]
FANWAI_TITLE = "番外篇"


def parse_chapter(fname):
    """返回 (是否番外, 章标题, 正文文本)。"""
    path = os.path.join(CHAPTERS, fname)
    with open(path, "r", encoding="utf-8") as f:
        lines = f.read().splitlines()
    title = fname
    body_start = 0
    for i, ln in enumerate(lines):
        if ln.startswith("#"):
            title = ln.lstrip("#").strip()
            body_start = i + 1
            break
    while body_start < len(lines) and not lines[body_start].strip():
        body_start += 1
    body = "\n".join(lines[body_start:]).strip()
    return fname.startswith("番外"), title, body


def make_cover():
    """根据底图生成带书名和作者署名的封面图。"""
    from PIL import Image, ImageDraw, ImageFont

    # 优先使用已有的真实底图；否则查找目录中唯一的 png
    bg_path = COVER_BG
    if not os.path.exists(bg_path):
        pngs = [n for n in os.listdir(COVER_DIR) if n.lower().endswith(".png")]
        if pngs:
            bg_path = os.path.join(COVER_DIR, sorted(pngs)[-1])
        else:
            raise FileNotFoundError("未找到封面底图，请先生成 cover_art/cover_bg.png")

    img = Image.open(bg_path).convert("RGB")

    # 去除 AI 生成图常见的黑边/暗边填充：构造阈值 mask 后取非黑边界框
    threshold = 25
    gray = img.convert("L")
    mask = gray.point(lambda x: 255 if x > threshold else 0, mode="1")
    bbox = mask.getbbox()
    if bbox:
        img = img.crop(bbox)

    # 统一为竖版封面尺寸 1024x1536，保持比例并居中裁剪
    target_w, target_h = 1024, 1536
    target_ratio = target_w / target_h
    img_ratio = img.width / img.height
    if img_ratio > target_ratio:
        # 图更宽，按高度缩放后裁宽度
        new_h = target_h
        new_w = int(new_h * img_ratio)
        img = img.resize((new_w, new_h), Image.Resampling.LANCZOS)
        left = (img.width - target_w) // 2
        img = img.crop((left, 0, left + target_w, target_h))
    else:
        # 图更高，按宽度缩放后裁高度
        new_w = target_w
        new_h = int(new_w / img_ratio)
        img = img.resize((new_w, new_h), Image.Resampling.LANCZOS)
        top = (img.height - target_h) // 2
        img = img.crop((0, top, target_w, top + target_h))

    draw = ImageDraw.Draw(img)

    # 字体
    font_title = ImageFont.truetype("C:/Windows/Fonts/msyhbd.ttc", 86, index=0)
    font_author = ImageFont.truetype("C:/Windows/Fonts/msyhbd.ttc", 44, index=0)
    font_tag = ImageFont.truetype("C:/Windows/Fonts/msyh.ttc", 28, index=0)

    def text_size(text, font):
        bbox = draw.textbbox((0, 0), text, font=font)
        return bbox[2] - bbox[0], bbox[3] - bbox[1]

    def draw_with_stroke(text, x, y, font, fill, stroke="black", width=3):
        for dx in range(-width, width + 1):
            for dy in range(-width, width + 1):
                if dx == 0 and dy == 0:
                    continue
                draw.text((x + dx, y + dy), text, font=font, fill=stroke)
        draw.text((x, y), text, font=font, fill=fill)

    # 书名：两行
    title_lines = ["百年回响", "次元归途"]
    y = 130
    for line in title_lines:
        w, _ = text_size(line, font_title)
        x = (target_w - w) // 2
        draw_with_stroke(line, x, y, font_title, "#F8F4E8", stroke="#1A1A1A", width=4)
        y += 105

    # 作者
    author_text = f"{AUTHOR}  著"
    w, h = text_size(author_text, font_author)
    x = (target_w - w) // 2
    draw_with_stroke(author_text, x, 1360, font_author, "#F5E6C8", stroke="#1A1A1A", width=3)

    # 小标签
    tag_text = "民国旧梦 × 次元重逢"
    w, _ = text_size(tag_text, font_tag)
    x = (target_w - w) // 2
    draw_with_stroke(tag_text, x, 1435, font_tag, "#D4E2E8", stroke="#1A1A1A", width=2)

    img.save(COVER_FINAL, "PNG", optimize=True)
    print("[ok] cover:", COVER_FINAL)
    return COVER_FINAL


def main():
    # 生成封面
    cover_path = make_cover()

    # 读取章节
    items = []
    for fname in ORDER:
        if not os.path.exists(os.path.join(CHAPTERS, fname)):
            print("[warn] missing:", fname)
            continue
        items.append(parse_chapter(fname))

    structure = []
    for vol_title, lo, hi in VOLUMES:
        chapters = []
        for is_fanwai, title, body in items:
            if is_fanwai:
                continue
            m = re.match(r"第(\d+)章", title)
            if m and lo <= int(m.group(1)) <= hi:
                chapters.append((title, body))
        structure.append((vol_title, chapters))
    fanwai_chs = [(t, b) for (f, t, b) in items if f]
    structure.append((FANWAI_TITLE, fanwai_chs))

    # ============ 1. 整合 Markdown ============
    rel_cover = "cover_art/cover_final.png"
    md_lines = [
        f"# {BOOK_TITLE}",
        "",
        f"![封面]({rel_cover})",
        "",
        f"**作者：{AUTHOR}**",
        "",
        SUBTITLE,
        "",
        "---",
        "",
        "## 目录",
        "",
    ]
    for idx, (vol_title, chapters) in enumerate(structure, start=1):
        md_lines.append(f"- [{vol_title}](#vol{idx})")
        for j, (title, _) in enumerate(chapters, start=1):
            md_lines.append(f"  - [{title}](#vol{idx}c{j})")
    md_lines.append("")
    md_lines.append("---")
    md_lines.append("")
    for idx, (vol_title, chapters) in enumerate(structure, start=1):
        md_lines.append(f"## {vol_title} {{#vol{idx}}}")
        md_lines.append("")
        for j, (title, body) in enumerate(chapters, start=1):
            md_lines.append(f"### {title} {{#vol{idx}c{j}}}")
            md_lines.append("")
            md_lines.append(body)
            md_lines.append("")
    md_lines.append("---")
    md_lines.append("")
    md_lines.append("—— 全书完 ——")
    md_text = "\n".join(md_lines).rstrip() + "\n"

    md_path = os.path.join(BASE, f"{BOOK_TITLE}（全本）.md")
    with open(md_path, "w", encoding="utf-8") as f:
        f.write(md_text)
    print("[ok] markdown:", md_path)

    # ============ 2. Word(.docx) ============
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
    sec.top_margin = sec.bottom_margin = Cm(2.5)
    sec.left_margin = sec.right_margin = Cm(2.8)

    def set_style_font(style, cn_font, en_font, size, bold=None, color=None):
        style.font.name = en_font
        style.font.size = Pt(size)
        if bold is not None:
            style.font.bold = bold
        if color is not None:
            style.font.color.rgb = color
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:eastAsia"), cn_font)

    set_style_font(doc.styles["Normal"], "宋体", "Times New Roman", 12)
    pf = doc.styles["Normal"].paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(0)
    pf.first_line_indent = Pt(24)

    h1 = doc.styles["Heading 1"]
    set_style_font(h1, "黑体", "Arial", 22, bold=True)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.first_line_indent = Pt(0)

    h2 = doc.styles["Heading 2"]
    set_style_font(h2, "黑体", "Arial", 16, bold=True)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(8)
    h2.paragraph_format.first_line_indent = Pt(0)

    h3 = doc.styles["Heading 3"]
    set_style_font(h3, "黑体", "Arial", 14, bold=True)
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(6)
    h3.paragraph_format.first_line_indent = Pt(0)

    # 页脚页码
    footer = sec.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run()
    fld1 = OxmlElement("w:fldChar"); fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
    fld2 = OxmlElement("w:fldChar"); fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld1); run._r.append(instr); run._r.append(fld2)

    # 封面页：顶部图片 + 书名 + 作者 + 副标题
    pic_para = doc.add_paragraph()
    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_run = pic_para.add_run()
    pic_run.add_picture(cover_path, width=Cm(15.0))

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title_para.add_run(BOOK_TITLE)
    r.font.name = "黑体"; r.font.size = Pt(28); r.font.bold = True
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    author_para = doc.add_paragraph()
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = author_para.add_run(f"作者：{AUTHOR}")
    r.font.name = "楷体"; r.font.size = Pt(14)
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "楷体")

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub_para.add_run(SUBTITLE)
    r.font.name = "楷体"; r.font.size = Pt(12)
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "楷体")

    doc.add_page_break()

    # 目录页
    doc.add_paragraph("目 录", style="Heading 1")
    tip = doc.add_paragraph()
    tip.alignment = WD_ALIGN_PARAGRAPH.LEFT
    rt = tip.add_run("（本目录为 Word 自动目录域：打开后全选 → 按 F9 或右键 → 更新域，即可生成带页码的完整目录。）")
    rt.font.size = Pt(10); rt.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    tp = doc.add_paragraph()
    run = tp.add_run()
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t"); placeholder.text = "【目录将在更新域后生成】"
    f3 = OxmlElement("w:fldChar"); f3.set(qn("w:fldCharType"), "end")
    run._r.append(f1); run._r.append(instr); run._r.append(f2); run._r.append(placeholder); run._r.append(f3)
    doc.add_page_break()

    INLINE = re.compile(r"(\*\*.+?\*\*|\*.+?\*)")

    def add_inline(paragraph, text):
        for part in INLINE.split(text):
            if not part:
                continue
            if part.startswith("**") and part.endswith("**"):
                r = paragraph.add_run(part[2:-2]); r.bold = True
            elif part.startswith("*") and part.endswith("*") and len(part) > 2:
                r = paragraph.add_run(part[1:-1]); r.italic = True
            else:
                paragraph.add_run(part)

    def add_body_block(text):
        for para_text in text.split("\n"):
            line = para_text.strip()
            if not line:
                continue
            p = doc.add_paragraph()
            add_inline(p, line)

    for vol_title, chapters in structure:
        h = doc.add_paragraph(vol_title, style="Heading 2")
        h.paragraph_format.page_break_before = True
        first_chapter_of_vol = True
        for title, body in chapters:
            h3p = doc.add_paragraph(title, style="Heading 3")
            if not first_chapter_of_vol:
                h3p.paragraph_format.page_break_before = True
            first_chapter_of_vol = False
            add_body_block(body)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.page_break_before = True
    r = p.add_run("—— 全书完 ——")
    r.font.size = Pt(14); r.font.bold = True

    docx_path = os.path.join(BASE, f"{BOOK_TITLE}（全本）.docx")
    doc.save(docx_path)
    print("[ok] docx:", docx_path)

    # ============ 3. EPUB ============
    import markdown as md_lib
    import ebooklib
    from ebooklib import epub
    from ebooklib.epub import EpubImage

    CSS = """
body { font-family: "Noto Serif SC", "Source Han Serif SC", "SimSun", serif;
       line-height: 1.7; margin: 5% 6%; }
h1 { font-size: 1.6em; text-align: center; margin: 0.8em 0 0.4em; }
h2 { font-size: 1.35em; margin: 1.2em 0 0.5em; page-break-before: always; }
h3 { font-size: 1.15em; margin: 1em 0 0.4em; page-break-before: always; }
p  { text-indent: 2em; margin: 0.35em 0; }
blockquote { margin: 0.6em 1.5em; color: #444; font-style: italic; }
hr { border: none; border-top: 1px solid #999; margin: 1em 4em; }
.center { text-align: center; }
.subtitle { text-align: center; font-size: 1.05em; color: #555; }
.author { text-align: center; font-size: 1.1em; margin: 0.6em 0; }
.theend { text-align: center; font-size: 1.1em; margin-top: 2em; }
.cover-img { width: 100%; max-width: 600px; display: block; margin: 0 auto 1em; }
"""

    def md_to_html(text):
        return md_lib.markdown(text, extensions=["extra"])

    book = epub.EpubBook()
    book.set_identifier("shiji-100-echo-2026")
    book.set_title(BOOK_TITLE)
    book.set_language("zh-CN")
    book.add_author(AUTHOR)

    # 封面图作为内嵌资源
    with open(cover_path, "rb") as cf:
        cover_img = EpubImage(
            uid="cover-image",
            file_name="images/cover.png",
            media_type="image/png",
            content=cf.read(),
        )
    book.add_item(cover_img)

    style_sheet = epub.EpubItem(
        uid="style", file_name="style/style.css", media_type="text/css", content=CSS
    )
    book.add_item(style_sheet)

    # 扉页/封面页
    title_page = epub.EpubHtml(title="封面", file_name="title.xhtml", lang="zh-CN")
    title_page.content = f"""
<img src="images/cover.png" alt="{BOOK_TITLE} 封面" class="cover-img"/>
<h1>{BOOK_TITLE}</h1>
<p class="author">作者：{AUTHOR}</p>
<p class="subtitle">{SUBTITLE}</p>
"""
    book.add_item(title_page)

    # 目录页
    toc_html = ["<h1>目录</h1><ul>"]
    toc_items = []
    for idx, (vol_title, chapters) in enumerate(structure, start=1):
        toc_html.append(f"<li><strong>{vol_title}</strong><ul>")
        vol_links = []
        for j, (title, body) in enumerate(chapters, start=1):
            fname = f"vol{idx}c{j}.xhtml"
            item = epub.EpubHtml(title=title, file_name=fname, lang="zh-CN")
            item.content = f'<h3>{title}</h3><div>{md_to_html(body)}</div>'
            item.add_item(style_sheet)
            book.add_item(item)
            vol_links.append(item)
            toc_html.append(f'<li><a href="{fname}">{title}</a></li>')
        toc_html.append("</ul></li>")
        toc_items.append((epub.Section(vol_title), vol_links))
    toc_html.append("</ul>")
    toc_page = epub.EpubHtml(title="目录", file_name="toc.xhtml", lang="zh-CN")
    toc_page.content = "".join(toc_html)
    book.add_item(toc_page)

    # 标记为 EPUB 封面元数据
    book.add_metadata(None, 'meta', '', {'name': 'cover', 'content': 'cover-image'})

    book.toc = toc_items
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    book.spine = ["nav", title_page, toc_page] + [
        item for _, links in toc_items for item in links
    ]

    epub_path = os.path.join(BASE, f"{BOOK_TITLE}（全本）.epub")
    epub.write_epub(epub_path, book, {})
    print("[ok] epub:", epub_path)

    # ============ 4. 统计 ============
    cn_chars = len(re.sub(r"\s", "", md_text))
    print(f"[info] 整合文本总字符数（不含空白）：{cn_chars}")
    print("[info] 卷数：%d；章节数：%d" % (len(structure), sum(len(c) for _, c in structure)))


if __name__ == "__main__":
    main()
